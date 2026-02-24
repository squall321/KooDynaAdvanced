"""
MODELING_TECHNICAL_DOCUMENT.md → .docx 변환 스크립트
====================================================
- pypandoc (pandoc): Markdown + LaTeX math → Word OMML 수식
- python-docx: 스타일 후처리 (폰트, 여백, 표 서식)

Usage:
    python export_docx.py
    python export_docx.py --input MY_DOC.md --output MY_DOC.docx
"""

import argparse
import logging
import os
import sys
from pathlib import Path

from battery_utils import setup_logger

logger = logging.getLogger(__name__)

DEFAULT_SRC = "MODELING_TECHNICAL_DOCUMENT.md"
DEFAULT_OUT = "MODELING_TECHNICAL_DOCUMENT.docx"


def convert_md_to_docx(src: str, out: str, log: logging.Logger) -> None:
    """Markdown → docx 변환 + 스타일 후처리"""
    try:
        import pypandoc
    except ImportError:
        log.error("pypandoc 미설치. pip install pypandoc-binary")
        sys.exit(1)
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from lxml import etree  # type: ignore[import-untyped]
    except ImportError:
        log.error("python-docx / lxml 미설치. pip install python-docx lxml")
        sys.exit(1)

    if not Path(src).exists():
        log.error("입력 파일 없음: %s", src)
        sys.exit(1)

    # ── 1. pandoc 변환 ──
    log.info("[1/3] pandoc 변환 중... %s → %s", src, out)
    pypandoc.convert_file(
        src, "docx",
        outputfile=out,
        extra_args=[
            "--standalone",
            "--toc",
            "--toc-depth=2",
            "--highlight-style=tango",
            "--wrap=none",
        ],
    )
    log.info("      → %s 생성 완료", out)

    # ── 2. python-docx 후처리 ──
    log.info("[2/3] 스타일 후처리 중...")
    doc = Document(out)

    # --- 페이지 설정 ---
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)

    # --- 기본 폰트 ---
    style = doc.styles["Normal"]
    font = style.font
    font.name = "맑은 고딕"
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    # CJK fallback
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn("w:rFonts"))
    rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    # --- 제목 스타일 ---
    heading_styles = {
        "Heading 1": (Pt(18), RGBColor(0x1B, 0x3A, 0x6B), True),
        "Heading 2": (Pt(14), RGBColor(0x2C, 0x5F, 0x2D), True),
        "Heading 3": (Pt(12), RGBColor(0x4A, 0x4A, 0x4A), True),
        "Heading 4": (Pt(11), RGBColor(0x5A, 0x5A, 0x5A), True),
    }
    for name, (size, color, bold) in heading_styles.items():
        if name in doc.styles:
            s = doc.styles[name]
            s.font.name = "맑은 고딕"
            s.font.size = size
            s.font.color.rgb = color
            s.font.bold = bold
            s.paragraph_format.space_before = Pt(14)
            s.paragraph_format.space_after = Pt(6)
            rPr2 = s.element.get_or_add_rPr()
            rF2 = rPr2.find(qn("w:rFonts"))
            if rF2 is None:
                rF2 = etree.SubElement(rPr2, qn("w:rFonts"))
            rF2.set(qn("w:eastAsia"), "맑은 고딕")

    # --- 표 서식 ---
    for table in doc.tables:  # pylint: disable=protected-access
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl = table._tbl  # pylint: disable=protected-access
        tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()  # pylint: disable=protected-access

        borders = tblPr.find(qn("w:tblBorders"))
        if borders is None:
            borders = etree.SubElement(tblPr, qn("w:tblBorders"))
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            elem = borders.find(qn(f"w:{edge}"))
            if elem is None:
                elem = etree.SubElement(borders, qn(f"w:{edge}"))
            elem.set(qn("w:val"), "single")
            elem.set(qn("w:sz"), "4")
            elem.set(qn("w:space"), "0")
            elem.set(qn("w:color"), "999999")

        if len(table.rows) > 0:
            for cell in table.rows[0].cells:
                tc = cell._tc  # pylint: disable=protected-access
                tcPr = tc.tcPr if tc.tcPr is not None else tc._add_tcPr()  # pylint: disable=protected-access
                shading = tcPr.find(qn("w:shd"))
                if shading is None:
                    shading = etree.SubElement(tcPr, qn("w:shd"))
                shading.set(qn("w:val"), "clear")
                shading.set(qn("w:color"), "auto")
                shading.set(qn("w:fill"), "1B3A6B")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.bold = True
                        run.font.size = Pt(9)

        for i, row in enumerate(table.rows):
            if i == 0:
                continue
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)
                    if i % 2 == 0:
                        tc = cell._tc  # pylint: disable=protected-access
                        tcPr = tc.tcPr if tc.tcPr is not None else tc._add_tcPr()  # pylint: disable=protected-access
                        shading = tcPr.find(qn("w:shd"))
                        if shading is None:
                            shading = etree.SubElement(tcPr, qn("w:shd"))
                        shading.set(qn("w:val"), "clear")
                        shading.set(qn("w:color"), "auto")
                        shading.set(qn("w:fill"), "F2F6FA")

    # --- 코드 블록 폰트 ---
    for para in doc.paragraphs:
        if para.style and para.style.name and "Source" in para.style.name:
            for run in para.runs:
                run.font.name = "Consolas"
                run.font.size = Pt(8.5)

    # --- 저장 ---
    doc.save(out)
    file_kb = os.path.getsize(out) / 1024
    log.info("[3/3] 후처리 완료 → %s (%.0f KB)", out, file_kb)


def main():
    parser = argparse.ArgumentParser(
        description="Markdown → Word (.docx) 변환 + 스타일 후처리")
    parser.add_argument("--input", "-i", type=str, default=DEFAULT_SRC,
                        help=f"입력 Markdown 파일 (default: {DEFAULT_SRC})")
    parser.add_argument("--output", "-o", type=str, default=DEFAULT_OUT,
                        help=f"출력 docx 파일 (default: {DEFAULT_OUT})")
    parser.add_argument("--verbose", "-v", action="store_true", help="상세 로그")
    args = parser.parse_args()

    log = setup_logger(
        "export_docx",
        level=logging.DEBUG if args.verbose else logging.INFO,
    )

    try:
        convert_md_to_docx(args.input, args.output, log)
    except (FileNotFoundError, OSError, ValueError) as e:
        log.error("변환 실패: %s", e, exc_info=True)
        sys.exit(1)


if __name__ == "__main__":
    main()
